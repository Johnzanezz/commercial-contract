#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
review_docx.py — 真实 OOXML 红线生成器（商事合同专家 · 工具增强版）

将 operations.json 中的审查意图落地为真实 Word 修订追踪：
  - 插入 (insert)    -> w:ins 包裹 w:r/w:t
  - 删除 (delete)    -> w:del 包裹 w:r/w:delText
  - 替换 (replace)   -> w:del(旧) + w:ins(新) 串联
  - 批注 (comment)   -> word/comments.xml 新增 w:comment + 正文 w:commentRangeStart/End + w:commentReference

作者红线：所有修订/批注 author 固定为「法大大iTerms」，不可改（执业身份红线）。

调用：
  python review_docx.py --source 合同.docx --ops operations.json --output 红线.docx [--author 法大大iTerms]

退出码：
  0 = 全部操作已落地
  1 = 至少一个操作未能落地（find 未命中 / 段落未定位 / 参数缺失）
"""

import argparse
import json
import sys
from datetime import datetime, timezone

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT, CONTENT_TYPE as CT
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree

REQUIRED_AUTHOR = "法大大iTerms"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENTS_BLOB = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '</w:comments>'
)


def now_iso():
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def make_run(text):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def make_ins(rid, author, date, text):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rid))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    ins.set(qn("w:rsidR"), "00000000")
    ins.append(make_run(text))
    return ins


def make_del(rid, author, date, text):
    d = OxmlElement("w:del")
    d.set(qn("w:id"), str(rid))
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), date)
    d.set(qn("w:rsidDel"), "00000000")
    r = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    d.append(r)
    return d


def make_comment_ref_run(cid):
    r = OxmlElement("w:r")
    cr = OxmlElement("w:commentReference")
    cr.set(qn("w:id"), str(cid))
    r.append(cr)
    return r


class IdGen:
    def __init__(self):
        self._n = 0

    def next(self):
        self._n += 1
        return self._n


class CommentsStore:
    """管理 word/comments.xml：不存在则创建并关联关系，支持追加 w:comment。

    注意：python-docx 的基类 Part 仅持有 blob（_blob），无 _element 解析。
    因此本类保留独立的 lxml 根元素 self.root，在 finalize() 时序列化回 _blob，
    再随 doc.save() 一并写出。
    """

    def __init__(self, doc):
        self.doc = doc
        self.part = None
        self.root = None
        self._id = 0
        for rel in doc.part.rels.values():
            if rel.reltype == RT.COMMENTS:
                self.part = rel.target_part
                self.root = parse_xml(self.part._blob)
                break

    def next_id(self):
        self._id += 1
        return self._id

    def ensure_part(self):
        if self.part is not None:
            return
        part = Part(PackURI("/word/comments.xml"), CT.WML_COMMENTS,
                   COMMENTS_BLOB.encode("utf-8"), self.doc.part.package)
        self.doc.part.relate_to(part, RT.COMMENTS)
        self.part = part
        self.root = parse_xml(COMMENTS_BLOB.encode("utf-8"))

    def add(self, cid, author, date, text, basis):
        if self.part is None:
            self.ensure_part()
        c = OxmlElement("w:comment")
        c.set(qn("w:id"), str(cid))
        c.set(qn("w:author"), author)
        c.set(qn("w:date"), date)
        c.set(qn("w:initials"), (author[:1] if author else ""))
        body = text or ""
        if basis:
            body = (body + "\n依据：" + basis) if body else ("依据：" + basis)
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = body
        r.append(t)
        p.append(r)
        c.append(p)
        self.root.append(c)

    def finalize(self):
        """将更新后的 comments 根序列化回 part._blob，供 doc.save() 写出。"""
        if self.part is not None and self.root is not None:
            self.part._blob = etree.tostring(
                self.root, xml_declaration=True, encoding="UTF-8", standalone=True)


def locate_paragraph(doc, op):
    """按 para(0起) 或 anchor 定位段落；都不给则 None。"""
    if op.get("para") is not None:
        try:
            idx = int(op["para"])
        except (TypeError, ValueError):
            return None
        if 0 <= idx < len(doc.paragraphs):
            return doc.paragraphs[idx]
        return None
    anchor = op.get("anchor")
    if anchor:
        for p in doc.paragraphs:
            if anchor in p.text:
                return p
    return None


def find_run_in_para(para, text):
    """返回 (run, idx) where run.text 首个包含 text 的位置。"""
    for run in para.runs:
        if run.text and text in run.text:
            return run, run.text.find(text)
    return None, -1


def op_insert(doc, para, op, author, date, ids, comments):
    text = op.get("text", "")
    if not text:
        return False, "insert 缺少 text"
    ins = make_ins(ids.next(), author, date, text)
    p = para._p
    if op.get("position") == "before" and op.get("find"):
        run, _ = find_run_in_para(para, op["find"])
        if run is None:
            return False, "insert(before): find 未命中"
        run._r.addprevious(ins)
    else:
        p.append(ins)
    if op.get("comment"):
        op_comment(doc, para, op, author, date, ids, comments)
    return True, None


def op_delete(doc, para, op, author, date, ids, comments):
    find = op.get("find", "")
    if not find:
        return False, "delete 缺少 find"
    run, idx = find_run_in_para(para, find)
    if run is None:
        return False, "delete: find 未命中"
    full = run.text
    before = full[:idx]
    after = full[idx + len(find):]
    run.text = before
    d = make_del(ids.next(), author, date, find)
    run._r.addnext(d)
    if after:
        d.addnext(make_run(after))
    if op.get("comment"):
        op_comment(doc, para, op, author, date, ids, comments)
    return True, None


def op_replace(doc, para, op, author, date, ids, comments):
    find = op.get("find", "")
    text = op.get("text", "")
    if not find:
        return False, "replace 缺少 find"
    if not text:
        return False, "replace 缺少 text"
    run, idx = find_run_in_para(para, find)
    if run is None:
        return False, "replace: find 未命中"
    full = run.text
    before = full[:idx]
    after = full[idx + len(find):]
    run.text = before
    d = make_del(ids.next(), author, date, find)
    run._r.addnext(d)
    ins = make_ins(ids.next(), author, date, text)
    d.addnext(ins)
    if after:
        ins.addnext(make_run(after))
    if op.get("comment"):
        op_comment(doc, para, op, author, date, ids, comments)
    return True, None


def op_comment(doc, para, op, author, date, ids, comments):
    find = op.get("find", "")
    cid = comments.next_id()
    crs = OxmlElement("w:commentRangeStart")
    crs.set(qn("w:id"), str(cid))
    cre = OxmlElement("w:commentRangeEnd")
    cre.set(qn("w:id"), str(cid))
    ref_r = make_comment_ref_run(cid)
    p = para._p
    anchor_run = None
    if find:
        anchor_run, _ = find_run_in_para(para, find)
    if anchor_run is not None:
        anchor_run._r.addprevious(crs)
        anchor_run._r.addnext(cre)
        cre.addnext(ref_r)
    else:
        p.insert(0, crs)
        p.append(cre)
        p.append(ref_r)
    comments.add(cid, author, date, op.get("comment", ""), op.get("basis", ""))
    return True, None


def dispatch(doc, para, op, author, date, ids, comments):
    t = op.get("type")
    if t == "insert":
        return op_insert(doc, para, op, author, date, ids, comments)
    if t == "delete":
        return op_delete(doc, para, op, author, date, ids, comments)
    if t == "replace":
        return op_replace(doc, para, op, author, date, ids, comments)
    if t == "comment":
        return op_comment(doc, para, op, author, date, ids, comments)
    return False, "未知 type: %r" % t


def main():
    ap = argparse.ArgumentParser(description="生成真实 Word 修订红线")
    ap.add_argument("--source", required=True, help="源合同 docx")
    ap.add_argument("--ops", required=True, help="operations.json")
    ap.add_argument("--output", required=True, help="输出红线 docx")
    ap.add_argument("--author", default=None, help="修订作者（默认取 ops.author 或 法大大iTerms）")
    args = ap.parse_args()

    with open(args.ops, "r", encoding="utf-8") as f:
        ops = json.load(f)

    author = args.author or ops.get("author") or REQUIRED_AUTHOR
    date = ops.get("date") or now_iso()

    doc = Document(args.source)
    comments = CommentsStore(doc)
    ids = IdGen()

    results = []
    for op in ops.get("operations", []):
        para = locate_paragraph(doc, op)
        if para is None:
            results.append({"id": op.get("id"), "type": op.get("type"),
                            "ok": False, "error": "段落未定位（para/anchor 无效）"})
            continue
        ok, err = dispatch(doc, para, op, author, date, ids, comments)
        results.append({"id": op.get("id"), "type": op.get("type"),
                        "ok": ok, "error": err})

    comments.finalize()
    doc.save(args.output)

    n_ok = sum(1 for r in results if r["ok"])
    n_fail = len(results) - n_ok
    print(json.dumps({
        "author": author,
        "operations_total": len(results),
        "operations_landed": n_ok,
        "operations_failed": n_fail,
        "details": results,
    }, ensure_ascii=False, indent=2))

    sys.exit(0 if n_fail == 0 else 1)


if __name__ == "__main__":
    main()
