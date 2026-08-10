#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
test_redline_pipeline.py — OOXML 红线引擎回归测试（commercial-contract）

固化端到端门禁语义，防止后续改动破坏：
  - 正向：4 类操作全落地，review 退出 0，validate 退出 0，verification_record 含 6 字段；
  - 负向 A：作者红线被突破 → validate 退出 2 (policy_blocked)；
  - 负向 B：操作未落地 (find 不存在) → review 退出 1，validate 退出 1 (needs_retry)。

自包含：在临时目录内生成夹具（合同.docx + operations.json），调用同目录的
review_docx.py / validate_review_outputs.py，断言退出码与产物结构。

用法：
  python test_redline_pipeline.py            # 用当前解释器（需 python-docx/lxml）
  python test_redline_pipeline.py --python /path/to/python.exe   # 指定解释器
"""

import argparse
import json
import os
import subprocess
import sys
import tempfile
import zipfile

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
REVIEW = os.path.join(HERE, "review_docx.py")
VALIDATE = os.path.join(HERE, "validate_review_outputs.py")

REQUIRED_AUTHOR = "法大大iTerms"
REQUIRED_VR_FIELDS = [
    "schema_version", "expert", "generated_at",
    "operations_summary", "disclaimer_checked", "policy_status",
]


def write_fixture(contract_path, ops_path):
    doc = Document()
    doc.add_paragraph("甲方：云南某建设工程有限公司")
    doc.add_paragraph("乙方：昆明某材料供应有限公司")
    doc.add_paragraph("第一条 标的与数量：乙方供应螺纹钢 1000 吨。")
    doc.add_paragraph("第二条 价款：合同总价暂定人民币伍佰万元整。")
    doc.add_paragraph("第三条 违约责任：任何一方违约应支付违约金每日千分之五。")
    doc.add_paragraph("第四条 争议解决：提交北京仲裁委员会仲裁。")
    doc.save(contract_path)

    ops = {
        "author": REQUIRED_AUTHOR,
        "date": "2026-08-06T09:30:00Z",
        "operations": [
            {"id": 1, "type": "insert", "para": 2, "text": "（含税）",
             "position": "append", "risk": "medium", "basis": "《民法典》第510条"},
            {"id": 2, "type": "delete", "para": 3, "find": "暂定",
             "risk": "high", "basis": "《民法典》第511条", "comment": "总价不应使用‘暂定’等不确定表述"},
            {"id": 3, "type": "replace", "para": 4, "find": "每日千分之五", "text": "每日万分之五",
             "risk": "high", "basis": "《民法典》第585条", "comment": "约定违约金过高，依司法解释可调减"},
            {"id": 4, "type": "comment", "para": 5, "find": "北京仲裁委员会",
             "risk": "high", "basis": "《仲裁法》第18条", "comment": "仲裁地点与甲方所在地不符，建议改为昆明仲裁机构"},
        ],
    }
    with open(ops_path, "w", encoding="utf-8") as f:
        json.dump(ops, f, ensure_ascii=False, indent=2)


def run(py, script, args):
    cmd = [py, script] + args
    p = subprocess.run(cmd, capture_output=True, text=True)
    return p.returncode, p.stdout, p.stderr


def tamper_author(redline_path, out_path):
    with zipfile.ZipFile(redline_path) as z:
        data = {n: z.read(n) for n in z.namelist()}
    doc = data["word/document.xml"].decode("utf-8")
    doc2 = doc.replace('w:author="%s"' % REQUIRED_AUTHOR, 'w:author="HACKER"', 1)
    data["word/document.xml"] = doc2.encode("utf-8")
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in data.items():
            z.writestr(n, b)


def check_json(path, must_have):
    with open(path, "r", encoding="utf-8") as f:
        obj = json.load(f)
    missing = [k for k in must_have if k not in obj]
    return obj, missing


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--python", default=sys.executable, help="运行脚本的解释器（需 python-docx/lxml）")
    args = ap.parse_args()
    py = args.python

    tmp = tempfile.mkdtemp(prefix="redline_regtest_")
    contract = os.path.join(tmp, "合同.docx")
    ops = os.path.join(tmp, "operations.json")
    redline = os.path.join(tmp, "红线.docx")
    redline_bad = os.path.join(tmp, "红线_篡改.docx")
    vr = os.path.join(tmp, "vr.json")
    rp = os.path.join(tmp, "rp.json")
    ops_bad = os.path.join(tmp, "ops_bad.json")

    write_fixture(contract, ops)

    results = []

    # ---- 正向：全落地 ----
    rc, out, err = run(py, REVIEW, ["--source", contract, "--ops", ops, "--output", redline])
    results.append(("review 正向 退出0", rc == 0, "rc=%d err=%s" % (rc, err[:200])))
    rc, out, err = run(py, VALIDATE,
                       ["--redline", redline, "--ops", ops, "--verification", vr, "--report", rp])
    results.append(("validate 正向 退出0", rc == 0, "rc=%d err=%s" % (rc, err[:200])))
    vr_obj, missing = check_json(vr, REQUIRED_VR_FIELDS)
    results.append(("verification_record 六字段齐全", not missing,
                    "missing=%s" % missing))
    results.append(("verification_record policy_status=passed",
                    vr_obj.get("policy_status") == "passed", str(vr_obj.get("policy_status"))))

    # 带 --disclaimer-checked 重跑，disclaimer_checked 应为 true
    rc, out, err = run(py, VALIDATE,
                       ["--redline", redline, "--ops", ops, "--verification", vr,
                        "--report", rp, "--disclaimer-checked"])
    vr_obj, _ = check_json(vr, REQUIRED_VR_FIELDS)
    results.append(("disclaimer_checked=true", vr_obj.get("disclaimer_checked") is True,
                    str(vr_obj.get("disclaimer_checked"))))

    # ---- 负向 A：作者红线被突破 → exit 2 ----
    tamper_author(redline, redline_bad)
    rc, out, err = run(py, VALIDATE,
                       ["--redline", redline_bad, "--ops", ops, "--verification", vr, "--report", rp])
    results.append(("validate 作者篡改 退出2(policy_blocked)", rc == 2, "rc=%d" % rc))

    # ---- 负向 B：操作未落地 → review 1, validate 1 ----
    with open(ops_bad, "w", encoding="utf-8") as f:
        json.dump({
            "author": REQUIRED_AUTHOR,
            "date": "2026-08-06T09:30:00Z",
            "operations": [
                {"id": 1, "type": "insert", "para": 2, "text": "（含税）"},
                {"id": 2, "type": "delete", "para": 3, "find": "不存在的文本XYZ",
                 "comment": "应删除但找不到"},
            ],
        }, f, ensure_ascii=False, indent=2)
    redline_bad2 = os.path.join(tmp, "红线_bad.docx")
    rc, out, err = run(py, REVIEW, ["--source", contract, "--ops", ops_bad, "--output", redline_bad2])
    results.append(("review 未落地 退出1", rc == 1, "rc=%d" % rc))
    rc, out, err = run(py, VALIDATE,
                       ["--redline", redline_bad2, "--ops", ops_bad,
                        "--verification", vr, "--report", rp])
    results.append(("validate 未落地 退出1(needs_retry)", rc == 1, "rc=%d" % rc))

    # ---- 汇总 ----
    print("=" * 60)
    print("OOXML 红线引擎回归测试")
    print("=" * 60)
    all_ok = True
    for name, ok, detail in results:
        mark = "PASS" if ok else "FAIL"
        if not ok:
            all_ok = False
        print("[%s] %s  (%s)" % (mark, name, detail))
    print("=" * 60)
    print("RESULT: %s" % ("ALL PASS" if all_ok else "HAS FAILURE"))
    print("临时目录: %s" % tmp)
    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
