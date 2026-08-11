#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
test_review_docx.py — review_docx.py 函数级单元测试

覆盖：
  - make_ins / make_del：修订元素结构与 author 红线属性
  - IdGen：单调递增
  - locate_paragraph：按 para 索引 / anchor / 无效输入
  - dispatch：按 type 路由至对应处理函数
  - op_insert / op_delete / op_replace / op_comment：落地语义（命中/未命中）
  - CommentsStore：缺失自动创建 + 追加 w:comment
  - 作者红线（最关键）：ops.author / --author 传入异值，生成产物 author 仍恒为 REQUIRED_AUTHOR
  - 端到端：subprocess 调用脚本，校验 w:ins/w:del/w:comment 落地与退出码

运行：python -m unittest tests.test_review_docx -v
"""

import importlib.util
import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile

from docx import Document

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SCRIPT = os.path.join(HERE, "review_docx.py")
VALIDATE = os.path.join(HERE, "validate_review_outputs.py")

_spec = importlib.util.spec_from_file_location("review_docx", SCRIPT)
rd = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(rd)

_vro_spec = importlib.util.spec_from_file_location("validate_review_outputs", VALIDATE)
vro = importlib.util.module_from_spec(_vro_spec)
_vro_spec.loader.exec_module(vro)


def build_contract(path):
    doc = Document()
    doc.add_paragraph("甲方：云南某建设工程有限公司")
    doc.add_paragraph("乙方：昆明某材料供应有限公司")
    doc.add_paragraph("第一条 标的与数量：乙方供应螺纹钢 1000 吨。")
    doc.add_paragraph("第二条 价款：合同总价暂定人民币伍佰万元整。")
    doc.add_paragraph("第三条 违约责任：任何一方违约应支付违约金每日千分之五。")
    doc.add_paragraph("第四条 争议解决：提交北京仲裁委员会仲裁。")
    doc.save(path)


class TestMakeRevisions(unittest.TestCase):
    def test_make_ins_structure(self):
        ins = rd.make_ins(1, rd.REQUIRED_AUTHOR, "2026-08-06T00:00:00Z", "（含税）")
        self.assertEqual(ins.get(rd.qn("w:author")), rd.REQUIRED_AUTHOR)
        self.assertEqual(ins.get(rd.qn("w:id")), "1")
        self.assertIn("（含税）", vro.ins_text(ins))

    def test_make_del_structure(self):
        d = rd.make_del(2, rd.REQUIRED_AUTHOR, "2026-08-06T00:00:00Z", "暂定")
        self.assertEqual(d.get(rd.qn("w:author")), rd.REQUIRED_AUTHOR)
        self.assertEqual(vro.del_text(d), "暂定")

    def test_idgen_monotonic(self):
        g = rd.IdGen()
        self.assertEqual([g.next() for _ in range(3)], [1, 2, 3])


class TestLocateParagraph(unittest.TestCase):
    def setUp(self):
        doc = Document()
        doc.add_paragraph("段落零")
        doc.add_paragraph("违约金每日千分之五条款")
        doc.add_paragraph("仲裁条款")
        self.doc = doc

    def test_by_index(self):
        p = rd.locate_paragraph(self.doc, {"para": 1})
        self.assertIsNotNone(p)
        self.assertEqual(p.text, "违约金每日千分之五条款")

    def test_by_anchor(self):
        p = rd.locate_paragraph(self.doc, {"anchor": "仲裁"})
        self.assertIsNotNone(p)
        self.assertEqual(p.text, "仲裁条款")

    def test_invalid_index(self):
        self.assertIsNone(rd.locate_paragraph(self.doc, {"para": 99}))
        self.assertIsNone(rd.locate_paragraph(self.doc, {"para": "abc"}))

    def test_no_hint(self):
        self.assertIsNone(rd.locate_paragraph(self.doc, {}))


class TestDispatch(unittest.TestCase):
    def setUp(self):
        doc = Document()
        doc.add_paragraph("原价款暂定伍佰万元。")
        self.doc = doc
        self.para = doc.paragraphs[0]
        self.ids = rd.IdGen()
        self.cs = rd.CommentsStore(doc)

    def test_dispatch_insert(self):
        ok, _ = rd.dispatch(self.doc, self.para,
                            {"type": "insert", "text": "（含税）"},
                            rd.REQUIRED_AUTHOR, "d", self.ids, self.cs)
        self.assertTrue(ok)
        self.assertTrue(any(rd.REQUIRED_AUTHOR == e.get(rd.qn("w:author"))
                            for e in self.para._p.findall(".//" + rd.qn("w:ins"))))

    def test_dispatch_unknown_type(self):
        ok, err = rd.dispatch(self.doc, self.para,
                              {"type": "bogus"},
                              rd.REQUIRED_AUTHOR, "d", self.ids, self.cs)
        self.assertFalse(ok)
        self.assertIn("未知", err)

    def test_op_delete_miss(self):
        ok, err = rd.op_delete(self.doc, self.para,
                               {"find": "不存在XYZ"},
                               rd.REQUIRED_AUTHOR, "d", self.ids, self.cs)
        self.assertFalse(ok)
        self.assertIn("未命中", err)

    def test_op_replace_lands_del_and_ins(self):
        ok, _ = rd.op_replace(self.doc, self.para,
                              {"find": "暂定", "text": "合计"},
                              rd.REQUIRED_AUTHOR, "d", self.ids, self.cs)
        self.assertTrue(ok)
        dels = self.para._p.findall(".//" + rd.qn("w:del"))
        ins = self.para._p.findall(".//" + rd.qn("w:ins"))
        self.assertTrue(any("暂定" in vro.del_text(e) for e in dels))
        self.assertTrue(any("合计" in vro.ins_text(e) for e in ins))


class TestCommentsStore(unittest.TestCase):
    def test_ensure_part_creates_comments(self):
        doc = Document()
        doc.add_paragraph("正文")
        cs = rd.CommentsStore(doc)
        self.assertIsNone(cs.part)
        cs.ensure_part()
        self.assertIsNotNone(cs.part)
        cs.add(1, rd.REQUIRED_AUTHOR, "d", "建议修改", "依据民法典")
        cs.finalize()
        root = cs.root
        comment = root.find(rd.qn("w:comment"))
        self.assertIsNotNone(comment)
        self.assertEqual(comment.get(rd.qn("w:author")), rd.REQUIRED_AUTHOR)
        body = "".join(t.text or "" for t in comment.iter(rd.qn("w:t")))
        self.assertIn("建议修改", body)
        self.assertIn("依据民法典", body)


class TestAuthorRedlineEnforced(unittest.TestCase):
    """关键：无论 ops.author / --author 请求何值，产物 author 恒为 REQUIRED_AUTHOR。"""

    def _run_review(self, ops_author=None, cli_author=None):
        tmp = tempfile.mkdtemp(prefix="rd_author_test_")
        contract = os.path.join(tmp, "c.docx")
        ops = os.path.join(tmp, "ops.json")
        out = os.path.join(tmp, "redline.docx")
        build_contract(contract)
        payload = {
            "author": ops_author or rd.REQUIRED_AUTHOR,
            "date": "2026-08-06T00:00:00Z",
            "operations": [
                {"id": 1, "type": "insert", "para": 2, "text": "（含税）"},
                {"id": 2, "type": "comment", "para": 5, "find": "北京仲裁委员会",
                 "comment": "建议改为昆明仲裁机构", "basis": "仲裁法第十八条"},
            ],
        }
        with open(ops, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        cmd = [sys.executable, SCRIPT, "--source", contract, "--ops", ops, "--output", out]
        if cli_author:
            cmd += ["--author", cli_author]
        p = subprocess.run(cmd, capture_output=True, text=True)
        return p, out

    def test_ops_author_hacker_ignored(self):
        p, out = self._run_review(ops_author="HACKER")
        self.assertEqual(p.returncode, 0, msg=p.stderr)
        combined = self._read_parts(out)
        self.assertNotIn("HACKER", combined)
        self.assertIn(rd.REQUIRED_AUTHOR, combined)

    def test_cli_author_hacker_ignored(self):
        p, out = self._run_review(cli_author="ATTACKER")
        self.assertEqual(p.returncode, 0, msg=p.stderr)
        combined = self._read_parts(out)
        self.assertNotIn("ATTACKER", combined)
        self.assertIn(rd.REQUIRED_AUTHOR, combined)

    def test_default_author_is_required(self):
        p, out = self._run_review()
        self.assertEqual(p.returncode, 0, msg=p.stderr)
        combined = self._read_parts(out)
        self.assertIn(rd.REQUIRED_AUTHOR, combined)

    @staticmethod
    def _read_parts(out):
        with zipfile.ZipFile(out) as z:
            combined = z.read("word/document.xml").decode("utf-8")
            if "word/comments.xml" in z.namelist():
                combined += z.read("word/comments.xml").decode("utf-8")
        return combined


class TestEndToEnd(unittest.TestCase):
    def test_four_ops_land_and_exit0(self):
        tmp = tempfile.mkdtemp(prefix="rd_e2e_")
        contract = os.path.join(tmp, "c.docx")
        ops = os.path.join(tmp, "ops.json")
        out = os.path.join(tmp, "redline.docx")
        build_contract(contract)
        payload = {
            "author": rd.REQUIRED_AUTHOR,
            "date": "2026-08-06T00:00:00Z",
            "operations": [
                {"id": 1, "type": "insert", "para": 2, "text": "（含税）"},
                {"id": 2, "type": "delete", "para": 3, "find": "暂定",
                 "comment": "总价不应使用暂定"},
                {"id": 3, "type": "replace", "para": 4, "find": "每日千分之五", "text": "每日万分之五",
                 "comment": "违约金过高可调减"},
                {"id": 4, "type": "comment", "para": 5, "find": "北京仲裁委员会",
                 "comment": "建议改为昆明仲裁机构"},
            ],
        }
        with open(ops, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        p = subprocess.run([sys.executable, SCRIPT, "--source", contract,
                            "--ops", ops, "--output", out],
                           capture_output=True, text=True)
        self.assertEqual(p.returncode, 0, msg=p.stderr)
        with zipfile.ZipFile(out) as z:
            doc_xml = z.read("word/document.xml").decode("utf-8")
            self.assertIn("w:ins", doc_xml)
            self.assertIn("w:del", doc_xml)
            self.assertIn("w:commentReference", doc_xml)
            self.assertIn("word/comments.xml", z.namelist())

    def test_unlanded_find_exits1(self):
        tmp = tempfile.mkdtemp(prefix="rd_unlanded_")
        contract = os.path.join(tmp, "c.docx")
        ops = os.path.join(tmp, "ops_bad.json")
        out = os.path.join(tmp, "redline_bad.docx")
        build_contract(contract)
        payload = {
            "author": rd.REQUIRED_AUTHOR,
            "date": "2026-08-06T00:00:00Z",
            "operations": [
                {"id": 1, "type": "insert", "para": 2, "text": "（含税）"},
                {"id": 2, "type": "delete", "para": 3, "find": "不存在的文本XYZ"},
            ],
        }
        with open(ops, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        p = subprocess.run([sys.executable, SCRIPT, "--source", contract,
                            "--ops", ops, "--output", out],
                           capture_output=True, text=True)
        self.assertEqual(p.returncode, 1, msg=p.stderr)


if __name__ == "__main__":
    unittest.main()
