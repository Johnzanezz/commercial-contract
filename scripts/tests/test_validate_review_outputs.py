#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
test_validate_review_outputs.py — validate_review_outputs.py 函数级单元测试

覆盖：
  - ins_text / del_text：从修订元素萃取文本
  - gather_comments：解析 comments 部件，author 红线检测
  - op_landed：insert/delete/replace/comment 命中与未命中语义
  - 退出码语义（subprocess 端到端）：
        0 = passed（全落地、作者合规）
        1 = needs_retry（操作未落地 / 批注结构缺失）
        2 = policy_blocked（作者红线被突破）

运行：python -m unittest tests.test_validate_review_outputs -v
"""

import importlib.util
import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile

from docx import Document

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REVIEW = os.path.join(HERE, "review_docx.py")
VALIDATE = os.path.join(HERE, "validate_review_outputs.py")

_spec = importlib.util.spec_from_file_location("validate_review_outputs", VALIDATE)
vro = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(vro)

# review 模块仅用于构造真实修订元素 / 生成夹具
_rd_spec = importlib.util.spec_from_file_location("review_docx", REVIEW)
rd = importlib.util.module_from_spec(_rd_spec)
_rd_spec.loader.exec_module(rd)

QD = rd.qn  # 复用同一命名空间前缀 helper


def build_contract(path):
    doc = Document()
    doc.add_paragraph("甲方：云南某建设工程有限公司")
    doc.add_paragraph("乙方：昆明某材料供应有限公司")
    doc.add_paragraph("第一条 标的与数量：乙方供应螺纹钢 1000 吨。")
    doc.add_paragraph("第二条 价款：合同总价暂定人民币伍佰万元整。")
    doc.add_paragraph("第三条 违约责任：任何一方违约应支付违约金每日千分之五。")
    doc.add_paragraph("第四条 争议解决：提交北京仲裁委员会仲裁。")
    doc.save(path)


def generate_redline(tmp):
    contract = os.path.join(tmp, "c.docx")
    ops = os.path.join(tmp, "ops.json")
    out = os.path.join(tmp, "redline.docx")
    build_contract(contract)
    payload = {
        "author": vro.REQUIRED_AUTHOR,
        "date": "2026-08-06T00:00:00Z",
        "operations": [
            {"id": 1, "type": "insert", "para": 2, "text": "（含税）"},
            {"id": 2, "type": "delete", "para": 3, "find": "暂定",
             "comment": "总价不应使用暂定"},
            {"id": 3, "type": "replace", "para": 4, "find": "每日千分之五", "text": "每日万分之五",
             "comment": "违约金过高可调减"},
            {"id": 4, "type": "comment", "para": 5, "find": "北京仲裁委员会",
             "comment": "建议改为昆明仲裁机构"},
        ],
    }
    with open(ops, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    p = subprocess.run([sys.executable, REVIEW, "--source", contract,
                        "--ops", ops, "--output", out],
                       capture_output=True, text=True)
    assert p.returncode == 0, p.stderr
    return ops, out


def tamper_author(redline_path, out_path):
    """将首处修订 author 改为 HACKER，模拟作者红线被突破（仅替换一处即可触发）。"""
    with zipfile.ZipFile(redline_path) as z:
        data = {n: z.read(n) for n in z.namelist()}
    doc = data["word/document.xml"].decode("utf-8")
    doc2 = doc.replace('w:author="%s"' % vro.REQUIRED_AUTHOR, 'w:author="HACKER"', 1)
    data["word/document.xml"] = doc2.encode("utf-8")
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in data.items():
            z.writestr(n, b)


class TestTextExtraction(unittest.TestCase):
    def test_ins_text(self):
        ins = rd.make_ins(1, vro.REQUIRED_AUTHOR, "d", "（含税）")
        self.assertEqual(vro.ins_text(ins), "（含税）")

    def test_del_text(self):
        d = rd.make_del(2, vro.REQUIRED_AUTHOR, "d", "暂定")
        self.assertEqual(vro.del_text(d), "暂定")


class TestGatherComments(unittest.TestCase):
    def test_gather_from_generated(self):
        tmp = tempfile.mkdtemp(prefix="vro_gc_")
        _ops, redline = generate_redline(tmp)
        doc = Document(redline)
        root, texts, ok = vro.gather_comments(doc)
        self.assertIsNotNone(root)
        self.assertTrue(ok)
        self.assertTrue(any("建议改为昆明仲裁机构" in t for t in texts.values()))

    def test_gather_bad_author(self):
        tmp = tempfile.mkdtemp(prefix="vro_gc_bad_")
        _ops, redline = generate_redline(tmp)
        bad = os.path.join(tmp, "bad.docx")
        tamper_comments_author(redline, bad)
        doc = Document(bad)
        _root, _texts, ok = vro.gather_comments(doc)
        self.assertFalse(ok)


def tamper_comments_author(redline_path, out_path):
    with zipfile.ZipFile(redline_path) as z:
        data = {n: z.read(n) for n in z.namelist()}
    cx = data["word/comments.xml"].decode("utf-8")
    cx2 = cx.replace('w:author="%s"' % vro.REQUIRED_AUTHOR, 'w:author="HACKER"', 1)
    data["word/comments.xml"] = cx2.encode("utf-8")
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in data.items():
            z.writestr(n, b)


class TestOpLanded(unittest.TestCase):
    def setUp(self):
        self.ins = [rd.make_ins(1, vro.REQUIRED_AUTHOR, "d", "（含税）")]
        self.dels = [rd.make_del(2, vro.REQUIRED_AUTHOR, "d", "暂定")]
        self.comments = {"1": "建议改为昆明仲裁机构"}

    def test_insert_hit(self):
        ok, _ = vro.op_landed({"type": "insert", "text": "（含税）"},
                              self.ins, self.dels, self.comments)
        self.assertTrue(ok)

    def test_insert_miss(self):
        ok, _ = vro.op_landed({"type": "insert", "text": "凭空出现"},
                              self.ins, self.dels, self.comments)
        self.assertFalse(ok)

    def test_delete_hit(self):
        ok, _ = vro.op_landed({"type": "delete", "find": "暂定"},
                              self.ins, self.dels, self.comments)
        self.assertTrue(ok)

    def test_replace_needs_both(self):
        ok, _ = vro.op_landed({"type": "replace", "find": "暂定", "text": "（含税）"},
                              self.ins, self.dels, self.comments)
        self.assertTrue(ok)

    def test_comment_hit(self):
        ok, _ = vro.op_landed({"type": "comment", "comment": "昆明仲裁机构"},
                              self.ins, self.dels, self.comments)
        self.assertTrue(ok)

    def test_comment_missing_part(self):
        ok, reason = vro.op_landed({"type": "comment", "comment": "x"},
                                   self.ins, self.dels, {})
        self.assertFalse(ok)
        self.assertIn("comments", reason)


class TestExitCodes(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix="vro_exit_")
        self.ops, self.redline = generate_redline(self.tmp)
        self.vr = os.path.join(self.tmp, "vr.json")
        self.rp = os.path.join(self.tmp, "rp.json")

    def _validate(self, redline=None, ops=None, extra=None):
        cmd = [sys.executable, VALIDATE,
               "--redline", redline or self.redline,
               "--ops", ops or self.ops,
               "--verification", self.vr, "--report", self.rp]
        if extra:
            cmd += extra
        return subprocess.run(cmd, capture_output=True, text=True)

    def test_passed_exit0(self):
        p = self._validate()
        self.assertEqual(p.returncode, 0, msg=p.stderr)
        with open(self.vr, encoding="utf-8") as f:
            vr = json.load(f)
        self.assertEqual(vr["policy_status"], "passed")
        self.assertEqual(vr["disclaimer_checked"], False)

    def test_disclaimer_checked_flag(self):
        self._validate(extra=["--disclaimer-checked"])
        with open(self.vr, encoding="utf-8") as f:
            vr = json.load(f)
        self.assertTrue(vr["disclaimer_checked"])

    def test_author_violation_exit2(self):
        bad = os.path.join(self.tmp, "redline_bad.docx")
        tamper_author(self.redline, bad)
        p = self._validate(redline=bad)
        self.assertEqual(p.returncode, 2, msg=p.stderr)
        with open(self.vr, encoding="utf-8") as f:
            vr = json.load(f)
        self.assertEqual(vr["policy_status"], "policy_blocked")

    def test_unlanded_exit1(self):
        ops_bad = os.path.join(self.tmp, "ops_bad.json")
        with open(ops_bad, "w", encoding="utf-8") as f:
            json.dump({
                "author": vro.REQUIRED_AUTHOR,
                "date": "2026-08-06T00:00:00Z",
                "operations": [
                    {"id": 1, "type": "insert", "para": 2, "text": "（含税）"},
                    {"id": 2, "type": "delete", "para": 3, "find": "不存在XYZ"},
                ],
            }, f, ensure_ascii=False, indent=2)
        p = self._validate(ops=ops_bad)
        self.assertEqual(p.returncode, 1, msg=p.stderr)
        with open(self.vr, encoding="utf-8") as f:
            vr = json.load(f)
        self.assertEqual(vr["policy_status"], "needs_retry")


if __name__ == "__main__":
    unittest.main()
